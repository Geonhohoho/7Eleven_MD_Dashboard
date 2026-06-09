from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "MD_신상품_초도발주_대시보드_사용_매뉴얼.docx"

FONT = "NanumSquare"
GREEN = "147D43"
DARK_GREEN = "0B5D32"
LIGHT_GREEN = "EAF6EF"
PALE_GREEN = "F5FBF7"
INK = "172033"
MUTED = "667085"
BORDER = "D9E2DD"
LIGHT_GRAY = "F5F7F8"
RED = "B42318"
RED_BG = "FEF3F2"
AMBER = "B54708"
AMBER_BG = "FFFAEB"
BLUE = "175CD3"
BLUE_BG = "EFF8FF"
WHITE = "FFFFFF"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph_border(paragraph, color=GREEN, size="14"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_text(doc, text, *, size=10.5, color=INK, bold=False, after=6,
             before=0, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_rich_paragraph(doc, segments, *, after=6, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    for text, bold, color in segments:
        run = p.add_run(text)
        set_run_font(run, size=10.5, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    return p


def add_callout(doc, title, body, *, fill=LIGHT_GREEN, title_color=DARK_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=fill, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=title_color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_status_table(doc):
    rows = [
        ("과발주", "출고율 < 50%", "초도 물량의 절반 미만만 7일 내 소화", RED_BG, RED),
        ("정상발주", "50% ≤ 출고율 ≤ 100%", "초도 물량 범위 안에서 적정하게 소화", LIGHT_GREEN, DARK_GREEN),
        ("과소발주", "출고율 > 100%", "초도 물량을 초과하는 출고 수요가 발생할 가능성", AMBER_BG, AMBER),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1800, 2300, 5260])
    set_table_borders(table)
    headers = ["판정", "기준", "업무상 의미"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=10, bold=True, color=WHITE)
    for label, criterion, meaning, fill, color in rows:
        cells = table.add_row().cells
        for i, text in enumerate((label, criterion, meaning)):
            set_cell_shading(cells[i], fill)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=9.5, bold=(i == 0), color=(color if i == 0 else INK))
    set_table_geometry(table, [1800, 2300, 5260])
    doc.add_paragraph()
    return table


def add_feature_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [1900, 3000, 4460])
    set_table_borders(table)
    for idx, text in enumerate(("화면 요소", "무엇을 보여주는가", "어떻게 활용하는가")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, bold=True, color=WHITE)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            if row_idx % 2 == 1:
                set_cell_shading(cells[idx], PALE_GREEN)
            p = cells[idx].paragraphs[0]
            r = p.add_run(text)
            set_run_font(r, size=9.2, bold=(idx == 0), color=INK)
    set_table_geometry(table, [1900, 3000, 4460])
    doc.add_paragraph()
    return table


def add_step(doc, number, title, body, check=None):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [860, 8500])
    set_table_borders(table, color=BORDER, size="5")
    left, right = table.rows[0].cells
    set_cell_shading(left, GREEN)
    set_cell_shading(right, PALE_GREEN)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"STEP\n{number}")
    set_run_font(r, size=10, bold=True, color=WHITE)
    p2 = right.paragraphs[0]
    r2 = p2.add_run(title)
    set_run_font(r2, size=10.5, bold=True, color=DARK_GREEN)
    p3 = right.add_paragraph()
    p3.paragraph_format.space_after = Pt(0 if not check else 3)
    p3.paragraph_format.line_spacing = 1.2
    r3 = p3.add_run(body)
    set_run_font(r3, size=9.6, color=INK)
    if check:
        p4 = right.add_paragraph()
        p4.paragraph_format.space_after = Pt(0)
        r4 = p4.add_run(f"확인 포인트: {check}")
        set_run_font(r4, size=9.2, bold=True, color=AMBER)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (17, GREEN, 18, 8),
        2: (13.5, DARK_GREEN, 14, 7),
        3: (11.5, INK, 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[list_style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Caption Small" not in styles:
        style = styles.add_style("Caption Small", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(9)
        style.font.color.rgb = rgb(MUTED)
        style.paragraph_format.space_after = Pt(4)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("MD 신상품 초도발주 대시보드 | 사용자 매뉴얼")
    set_run_font(hr, size=8.5, bold=True, color=MUTED)
    set_paragraph_border(hp, color=BORDER, size="5")

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def build_document():
    doc = Document()
    style_document(doc)

    # Cover
    add_text(doc, "SEVEN-ELEVEN MD OPERATIONS", size=10, bold=True, color=GREEN, after=42)
    add_text(doc, "MD 신상품 초도발주\n대시보드 사용 매뉴얼",
             size=28, bold=True, color=INK, after=12)
    add_text(doc, "처음 사용하는 MD를 위한 목적·지표·업무 절차 안내서",
             size=14, color=MUTED, after=32)
    cover_rule = doc.add_paragraph()
    set_paragraph_border(cover_rule, color=GREEN, size="28")
    cover_rule.paragraph_format.space_after = Pt(28)

    add_callout(
        doc,
        "이 문서의 목표",
        "대시보드를 처음 보는 사용자가 별도의 구두 설명 없이도 금일 발주 대상 상품을 확인하고, "
        "발주 수량을 검토·조정·확정하며, 과거 상품의 실제 성과를 분석할 수 있도록 안내합니다.",
    )
    add_text(doc, "대상 사용자", size=10, bold=True, color=GREEN, after=3, before=12)
    add_text(doc, "신상품 초도발주를 검토하는 MD, 발주 결과를 분석하는 운영 담당자, 대시보드를 처음 인수받은 프로젝트 구성원",
             size=10.5, after=12)
    add_text(doc, "문서 기준", size=10, bold=True, color=GREEN, after=3)
    add_text(doc, "2026년 6월 9일 현재 구현된 React 대시보드 및 데이터 산출 로직 기준",
             size=10.5, after=12)
    add_text(doc, "중요한 전제", size=10, bold=True, color=GREEN, after=3)
    add_text(doc, "본 대시보드는 MD 의사결정을 지원하는 도구입니다. 권장값과 경고 색상은 검토의 출발점이며, "
             "프로모션·진열·공급 제약·상품 특성 등 현업 판단을 대체하지 않습니다.",
             size=10.5, after=18)
    add_text(doc, "Version 1.0", size=9, color=MUTED, after=0)

    add_page_break(doc)

    # How to read
    add_heading(doc, "0. 가장 먼저 알아둘 것", 1)
    add_callout(
        doc,
        "5분 안에 이해하는 화면 구조",
        "상단에는 두 개의 탭이 있습니다. ‘금일 신상품 작업’은 오늘 처리할 발주 업무를 수행하는 화면이고, "
        "‘과거 신상품 조회’는 이미 출시된 상품의 실제 발주·출고 결과를 분석하는 화면입니다.",
        fill=BLUE_BG,
        title_color=BLUE,
    )
    add_heading(doc, "두 탭의 역할", 2)
    add_feature_table(doc, [
        ("금일 신상품 작업", "선택 날짜에 마감되는 신상품과 ML 권장 발주량", "발주 박스 수를 입력하고 센터 분배를 검토한 뒤 발주를 확정"),
        ("과거 신상품 조회", "과거 상품의 실제 초도 발주량, 7일 출고율, 예약주문 및 센터별 성과", "유사 상품 탐색, 과거 의사결정 평가, 향후 발주 판단 근거 확보"),
    ])

    add_heading(doc, "초보 사용자의 기본 동선", 2)
    for text in (
        "상단 탭에서 수행하려는 업무를 선택합니다.",
        "날짜와 필터를 설정해 대상 상품을 좁힙니다.",
        "상품 행의 수량·출고율·색상 판정을 먼저 확인합니다.",
        "상품 행을 클릭하여 상세 그래프와 센터 정보를 확인합니다.",
        "금일 작업에서는 박스 수량을 입력하고 필요하면 센터 분배량·가중치를 조정합니다.",
        "최종 검토 후 ‘발주 확정’을 누릅니다.",
        "과거 조회에서는 결과를 비교하고 필요하면 엑셀로 내려받습니다.",
    ):
        add_number(doc, text)

    add_heading(doc, "색상은 무엇을 의미하나", 2)
    add_status_table(doc)
    add_callout(
        doc,
        "판정 용어 주의",
        "화면의 ‘과소발주’는 주문량이 부족해 결품 위험이 높아질 수 있다는 의미입니다. "
        "‘과발주’는 주문량이 수요보다 많아 재고가 남을 가능성이 높다는 의미입니다.",
        fill=AMBER_BG,
        title_color=AMBER,
    )

    # Concepts
    add_heading(doc, "1. 핵심 용어와 계산 기준", 1)
    add_heading(doc, "박스와 EA", 2)
    add_rich_paragraph(doc, [
        ("박스", True, GREEN),
        ("는 MD가 실제 입력하는 발주 단위입니다. ", False, INK),
        ("EA", True, GREEN),
        ("는 낱개 수량입니다. 상품별 입수량(LDU)에 따라 ‘박스 × 박스당 EA’로 환산됩니다.", False, INK),
    ])
    add_callout(
        doc,
        "예시",
        "1박스에 20EA가 들어 있는 상품을 100박스 발주하면 총 발주량은 2,000EA입니다. "
        "화면의 입력칸에는 100을 입력하고, 입력칸 아래에서 EA 환산값 2,000을 확인합니다.",
    )

    add_heading(doc, "권장 초도 발주량", 2)
    add_text(doc, "ML 예측 결과를 기반으로 제안되는 전체 초도 발주량입니다. 화면에서는 업무 입력 단위에 맞춰 박스로 표시하고, 상세 영역에서는 EA 기준 값도 함께 확인할 수 있습니다.")
    add_bullet(doc, "권장값은 자동 확정값이 아니라 비교 기준입니다.")
    add_bullet(doc, "사용자가 발주량을 바꾸면 예상 출고율과 센터별 분배량이 함께 변합니다.")
    add_bullet(doc, "입수량이 상품마다 다르므로 서로 다른 상품의 박스 수만 단순 비교하지 않습니다.")

    add_heading(doc, "금일 탭의 예상 출고율", 2)
    add_callout(
        doc,
        "현재 구현 기준",
        "금일 탭의 예상 출고율은 ‘ML 권장 발주 EA ÷ 사용자가 입력한 발주 EA × 100’으로 계산됩니다. "
        "따라서 권장량과 동일하게 입력하면 약 100%, 권장량보다 많이 입력하면 100%보다 낮아지고, "
        "적게 입력하면 100%보다 높아집니다.",
        fill=RED_BG,
        title_color=RED,
    )
    add_text(doc, "이 값은 실제 미래 출고량을 직접 예측한 판매율이 아니라, 권장량 대비 입력량의 적정성을 보여주는 현재의 통제 지표입니다. "
             "따라서 프로모션이나 공급 이슈까지 반영된 확정 수요예측으로 해석하면 안 됩니다.")

    add_heading(doc, "과거 탭의 실제 초도 발주량과 7일 출고율", 2)
    add_bullet(doc, "실제 초도 발주량: 센터 재고 데이터에서 출시일까지 발생한 양(+)의 재고 증가분을 센터별로 합산한 값")
    add_bullet(doc, "출시 후 7일 출고량: 출시일 재고와 출시 7일 후 재고의 순감소분을 기준으로 집계")
    add_bullet(doc, "실제 출고율: 출시 후 7일 출고량 ÷ 실제 초도 발주량 × 100")
    add_text(doc, "과거 탭의 출고율은 실제 재고 흐름을 이용한 사후 성과 지표입니다. 금일 탭의 예상 출고율과 계산 목적이 다릅니다.")

    add_heading(doc, "센터 분배 기준", 2)
    add_text(doc, "전체 발주 박스는 센터별 점포 비중과 가중치를 함께 반영하여 배분합니다.")
    add_callout(
        doc,
        "센터 배분 지수",
        "센터별 배분 지수 = 해당 센터 점포 비중 × 센터 가중치\n"
        "센터별 박스 수 = 전체 입력 박스 × (해당 센터 배분 지수 ÷ 전체 센터 배분 지수 합)",
        fill=BLUE_BG,
        title_color=BLUE,
    )
    add_text(doc, "반올림으로 발생하는 잔여 박스는 소수점 비중이 큰 센터부터 배정되어, 센터별 박스 합이 사용자가 입력한 전체 박스 수와 일치하도록 처리됩니다.")

    # Today tab
    add_heading(doc, "2. 금일 신상품 작업 탭", 1)
    add_callout(
        doc,
        "탭의 목적",
        "선택한 날짜에 발주 마감 대상이 되는 신상품을 확인하고, ML 권장량을 기준으로 실제 발주량을 검토·조정·확정하는 업무 화면입니다.",
    )

    add_heading(doc, "어떤 상품이 표시되는가", 2)
    add_text(doc, "현재 날짜 설정에서 선택한 날짜와 발주 마감일이 일치하고, 예약주문 4일 데이터가 모두 존재하며 각 일자의 예약주문 수량이 0보다 큰 상품만 표시됩니다.")
    add_bullet(doc, "날짜를 바꿨는데 상품이 없다면 해당 날짜에 마감되는 상품이 없거나 예약주문 4일 데이터가 완전하지 않을 수 있습니다.")
    add_bullet(doc, "날짜를 변경하면 화면의 발주 확정 상태는 초기화됩니다.")

    add_heading(doc, "상단 요약 카드", 2)
    add_feature_table(doc, [
        ("이번 주 신규 상품 수", "선택한 날짜에 조회된 고유 신상품 수", "오늘 처리해야 할 전체 업무량 파악"),
        ("검토 필요 상품 수", "예상 출고율이 50% 미만 또는 100% 초과인 상품 수", "과발주·결품 위험 상품을 우선 검토"),
        ("발주 완료", "현재 화면에서 발주 확정한 상품 수 / 전체 상품 수", "업무 완료 여부 확인"),
    ])

    add_heading(doc, "상품 목록의 각 열", 2)
    add_feature_table(doc, [
        ("상품코드 / 상품명", "대상 상품 식별 정보", "다른 자료와 대조하거나 상품 검색 시 사용"),
        ("카테고리", "상품 분류", "유사 상품과 특성을 판단하는 참고 정보"),
        ("판매가", "소비자 판매가격", "가격대와 수요 특성 검토"),
        ("예상 출고율", "현재 입력량 대비 ML 권장량 비율", "과발주·정상·과소발주 가능성 판단"),
        ("권장 발주량", "ML 권장 초도 발주량을 박스로 환산한 값", "입력량의 기본값 및 비교 기준"),
        ("발주 수량 입력", "사용자가 확정하려는 전체 박스 수", "직접 수정 가능하며 아래 EA 환산값 확인"),
        ("처리", "발주 확정 / 확정 완료", "상품별 작업 상태 기록"),
    ])

    add_heading(doc, "금일 작업: 단계별 사용 방법", 2)
    add_step(doc, 1, "현재 날짜 설정", "상단 달력에서 업무 기준일을 선택합니다.",
             "조회 가능 구간 안의 날짜인지, 오늘 실제 처리하려는 마감일인지 확인합니다.")
    add_step(doc, 2, "요약 카드 확인", "신규 상품 수, 검토 필요 상품 수, 발주 완료 수를 확인합니다.",
             "검토 필요 상품을 먼저 처리하면 위험 상품을 놓칠 가능성이 줄어듭니다.")
    add_step(doc, 3, "상품 행 1차 검토", "판매가, 예상 출고율, 권장 발주량을 확인합니다.",
             "빨강·주황 색상 상품은 원인을 확인하기 전 바로 확정하지 않습니다.")
    add_step(doc, 4, "상품 상세 열기", "상품 행을 한 번 클릭하면 바로 아래에 상세 정보가 열립니다. 다시 클릭하면 닫힙니다.",
             "입력칸이나 버튼을 누른 동작과 행 펼치기를 구분합니다.")
    add_step(doc, 5, "발주 박스 수 입력", "기본값은 권장 박스 수입니다. 기존 값을 지우고 원하는 숫자를 입력할 수 있습니다.",
             "입력칸 아래의 EA 환산값과 권장 대비 과다·과소 문구를 함께 봅니다.")
    add_step(doc, 6, "센터별 분배 검토", "전체 박스 수가 센터별로 어떻게 배분되는지 확인합니다. 센터별 입력칸을 직접 수정할 수도 있습니다.",
             "센터 수량을 수정하면 전체 발주 박스 합도 함께 변경됩니다.")
    add_step(doc, 7, "센터 가중치 조정", "점포 비중 외에 지역·상권 특성을 반영할 필요가 있으면 센터 가중치를 수정합니다.",
             "가중치 변경 후 센터별 분배량이 의도한 방향으로 바뀌었는지 확인합니다.")
    add_step(doc, 8, "발주 확정", "전체 수량과 센터 분배를 최종 검토한 뒤 ‘발주 확정’을 누릅니다.",
             "확정 후에도 버튼을 다시 눌러 확정 상태를 해제할 수 있습니다.")

    add_heading(doc, "상품 상세 화면 읽는 법", 2)
    add_feature_table(doc, [
        ("예약주문 4일 추세선", "발주 마감 기준 4일간의 일자별 예약주문", "초기 수요 강도와 증가·감소 패턴 확인"),
        ("예약주문 4일 총합", "4일 예약주문 합계", "ML 권장량과 함께 수요 규모를 판단"),
        ("유사상품 과거 발주량 조회", "같은 카테고리의 설정일 이전 상품을 최신순으로 표시", "신상품 판단 시 과거 사례를 보조 근거로 사용"),
        ("센터 가중치 수정", "센터별 점포 비중과 가중치", "지역별 차이를 배분량에 반영"),
        ("센터별 분배 수량", "전체 발주량을 센터별 박스로 배분", "센터별 과다·부족 여부를 검토하고 직접 수정"),
    ])

    add_callout(
        doc,
        "확정 전 필수 확인",
        "권장량을 그대로 사용하는 경우에도 예약주문 추세, EA 환산, 센터별 분배 합계, 가중치 변경 여부를 반드시 확인하십시오. "
        "권장량은 판단을 돕는 기준이지 자동 승인값이 아닙니다.",
        fill=AMBER_BG,
        title_color=AMBER,
    )

    # Past tab
    add_heading(doc, "3. 과거 신상품 조회 탭", 1)
    add_callout(
        doc,
        "탭의 목적",
        "과거 출시 상품의 실제 초도 발주량과 출시 후 7일 소화 성과를 조회하여, 과거 의사결정이 적정했는지 검토하고 향후 유사 상품 발주에 활용하는 분석 화면입니다.",
    )

    add_heading(doc, "필터 사용 방법", 2)
    add_feature_table(doc, [
        ("날짜", "출시일 시작~종료 구간", "특정 시즌·행사·분기의 상품만 조회"),
        ("대분류", "가장 넓은 상품 분류", "선택하면 중분류 선택지가 해당 대분류에 맞게 변경"),
        ("중분류", "대분류 아래의 세부 분류", "선택하면 소분류 선택지가 해당 중분류에 맞게 변경"),
        ("소분류", "가장 구체적인 상품 분류", "유사성이 높은 상품만 좁혀 조회"),
        ("정렬", "최신순·출고율 높은순·낮은순·발주량 높은순", "분석 목적에 맞춰 우선순위 변경"),
        ("검색", "상품명 또는 상품코드", "특정 상품을 빠르게 찾기"),
        ("엑셀 다운로드", "현재 필터 결과 전체", "추가 분석·보고 자료 작성에 활용"),
    ])
    add_text(doc, "분류 필터는 대분류 → 중분류 → 소분류 순으로 선택하는 것이 가장 안전합니다. 상위 분류를 바꾸면 하위 선택지를 다시 확인합니다.")

    add_heading(doc, "상품 목록 해석", 2)
    add_feature_table(doc, [
        ("상품코드 / 상품명 / 출시일", "상품 식별과 출시 시점", "동일 상품 및 시즌 비교"),
        ("대·중·소분류", "상품 분류 체계", "유사 상품군 구성"),
        ("과거 발주량", "센터별 실제 초도 배분량의 합", "당시 전체 초도 물량 확인"),
        ("출고율", "출시 후 7일 출고량 ÷ 실제 초도 발주량", "초도 물량의 7일 소화 성과 평가"),
    ])
    add_text(doc, "과거 발주량 숫자의 색상은 출고율 판정과 연동됩니다. 빨강은 과발주, 초록은 정상발주, 주황은 과소발주를 의미합니다.")
    add_callout(
        doc,
        "상세 조회가 열리지 않는 상품",
        "과거 발주량이 0인 상품은 상세 정보가 열리지 않습니다. 해당 행의 발주량 또는 출고율 아래에 센터 매핑 없음, "
        "출시일까지 재고 증가분 미확인 등 사유가 표시될 수 있습니다.",
        fill=AMBER_BG,
        title_color=AMBER,
    )

    add_heading(doc, "과거 상품 상세 화면 읽는 법", 2)
    add_heading(doc, "예약주문 수량 시각화", 3)
    add_bullet(doc, "실선: 날짜별 예약주문 수량")
    add_bullet(doc, "점선과 연한 영역: 예약주문 누적합")
    add_bullet(doc, "예약주문 초기 4일 합: 예약주문 접수 초기 4일간의 합계")
    add_bullet(doc, "예약주문 수량 합: 시각화된 전체 예약주문 기간의 합계")
    add_text(doc, "초기 반응과 전체 누적 반응을 함께 보면, 초반에 수요가 집중된 상품인지 뒤늦게 수요가 붙은 상품인지 구분할 수 있습니다.")

    add_heading(doc, "출시 후 7일 출고 흐름", 3)
    add_bullet(doc, "실선: 출시 다음 날부터 7일간의 일자별 센터 순출고량")
    add_bullet(doc, "점선과 연한 영역: 7일 누적 출고량")
    add_bullet(doc, "출시 후 센터 7일 출고 합: 센터별 출고량을 일자별로 합산한 값")
    add_text(doc, "예약주문 그래프와 동일한 축 기준을 사용하므로 출시 전 관심과 출시 후 실제 소화 규모를 시각적으로 비교할 수 있습니다.")

    add_heading(doc, "실제 센터 분배량·성과 맵", 3)
    add_bullet(doc, "회색 바: 센터에 배분된 실제 초도 물량")
    add_bullet(doc, "색상 바: 해당 초도 물량 중 출시 후 7일간 소화된 비율")
    add_bullet(doc, "행 위에 마우스를 올리면 센터의 7일 소화 물량이 즉시 표시")
    add_bullet(doc, "센터별 과발주·정상발주·과소발주 문구와 색상으로 판정 확인")
    add_text(doc, "센터별 분배량의 합은 상품 행의 과거 발주량과 일치합니다. 특정 센터만 과발주 또는 과소발주라면 전체 수량보다 배분 방식이 문제였을 가능성을 검토합니다.")

    add_heading(doc, "ML 기반 예측 비교", 3)
    add_feature_table(doc, [
        ("비교 상태", "ML 권장량 대비 실제 초도량 차이", "과대 발주·적정 발주·과소 발주 여부 확인"),
        ("ML 권장 초도", "모델이 제안한 초도 EA", "모델 기준점"),
        ("실제 초도", "센터 재고로 산출한 실제 초도 EA", "실행된 의사결정"),
        ("편차", "실제 초도 - ML 권장 초도", "수량 및 비율 차이 확인"),
        ("실제 출고율", "7일 실제 출고량 ÷ 실제 초도", "실행 결과 평가"),
        ("ML 기준 기대 출고율", "7일 실제 출고량 ÷ ML 권장량", "모델 권장량을 썼을 때의 사후 비교 지표"),
        ("정상 발주량 범주", "7일 출고량~7일 출고량의 2배", "출고율 50~100%를 만족하는 초도 수량 범위"),
    ])
    add_callout(
        doc,
        "ML 비교 해석 시 주의",
        "ML 권장량과 실제 발주량이 다르다는 사실만으로 어느 쪽이 맞았다고 단정할 수 없습니다. 실제 출고율, 센터별 편차, "
        "프로모션, 품절 여부, 추가 입고 여부를 함께 확인해야 합니다.",
        fill=BLUE_BG,
        title_color=BLUE,
    )

    add_heading(doc, "과거 조회: 단계별 사용 방법", 2)
    add_step(doc, 1, "분석 목적 정하기", "특정 상품 확인, 카테고리 벤치마킹, 과발주 사례 탐색 등 목적을 먼저 정합니다.")
    add_step(doc, 2, "날짜 범위 설정", "비교하려는 출시 기간을 선택합니다.", "시즌성이 강한 상품은 유사 시즌끼리 비교합니다.")
    add_step(doc, 3, "분류 필터 적용", "대분류 → 중분류 → 소분류 순으로 범위를 좁힙니다.")
    add_step(doc, 4, "정렬과 검색", "출고율 또는 발주량 정렬을 사용하거나 상품명·코드를 검색합니다.")
    add_step(doc, 5, "행 판정 확인", "과거 발주량의 색상과 출고율을 보고 우선 분석할 상품을 고릅니다.")
    add_step(doc, 6, "상품 상세 열기", "행을 클릭해 예약주문, 7일 출고, 센터 성과, ML 비교를 확인합니다.")
    add_step(doc, 7, "원인 가설 세우기", "전체 과발주인지, 특정 센터 배분 문제인지, ML 권장과 실제 실행의 차이인지 구분합니다.")
    add_step(doc, 8, "결과 활용", "유사 상품의 신규 발주 근거로 활용하거나, 필터 결과를 엑셀로 내려받아 보고 자료로 사용합니다.")

    add_heading(doc, "추천 분석 질문", 2)
    for question in (
        "예약주문 초기 4일 반응과 전체 예약주문 반응은 어떻게 달랐는가?",
        "실제 초도 물량 중 7일 내 몇 퍼센트가 소화되었는가?",
        "전체는 정상인데 특정 센터만 과발주 또는 과소발주였는가?",
        "ML 권장량과 실제 발주량의 차이는 얼마였고 결과는 어땠는가?",
        "동일 소분류 상품 중 반복적으로 과발주되는 가격대나 센터가 있는가?",
    ):
        add_bullet(doc, question)

    # Scenarios
    add_heading(doc, "4. 실제 업무 활용 시나리오", 1)
    add_heading(doc, "시나리오 A: 권장량보다 많이 발주하려는 경우", 2)
    add_text(doc, "ML 권장량이 1,000박스인데 프로모션을 고려해 1,200박스를 입력했다고 가정합니다.")
    add_bullet(doc, "입력량이 권장량보다 많으므로 예상 출고율은 100%보다 낮아집니다.")
    add_bullet(doc, "입력칸에 권장 대비 과다 발주 문구가 표시될 수 있습니다.")
    add_bullet(doc, "센터별 분배량이 자동으로 확대되므로, 수요가 약한 센터까지 과도하게 늘어나지 않았는지 확인합니다.")
    add_bullet(doc, "증량 근거를 프로모션, 점포 확대, 유사 상품 성과 등으로 설명할 수 있을 때 확정합니다.")

    add_heading(doc, "시나리오 B: 권장량보다 적게 발주하려는 경우", 2)
    add_text(doc, "ML 권장량이 1,000박스인데 공급 제약으로 800박스를 입력했다고 가정합니다.")
    add_bullet(doc, "예상 출고율은 100%를 초과하며 결품 위험 또는 과소발주 신호가 나타납니다.")
    add_bullet(doc, "센터별 최소 공급 필요량과 핵심 상권 센터를 우선 확인합니다.")
    add_bullet(doc, "센터 가중치를 조정하면 제한된 물량을 중요 센터에 집중할 수 있습니다.")
    add_bullet(doc, "공급 제약이 해소될 가능성과 추가 입고 계획을 함께 검토합니다.")

    add_heading(doc, "시나리오 C: 과거 상품의 전체 출고율은 정상이나 일부 센터가 과발주인 경우", 2)
    add_text(doc, "상품 전체는 70% 출고율로 정상 범주지만, 특정 센터가 30%라면 전체 수량보다 센터 배분이 문제였을 가능성이 있습니다.")
    add_bullet(doc, "해당 센터의 실제 초도 배분량과 7일 소화 물량을 마우스 오버로 확인합니다.")
    add_bullet(doc, "비슷한 센터의 점포 비중·가중치와 비교합니다.")
    add_bullet(doc, "향후 유사 상품에서 해당 센터의 가중치를 낮출 근거로 활용합니다.")

    add_heading(doc, "시나리오 D: 실제 초도가 ML 권장보다 크게 적었는데 출고율이 높은 경우", 2)
    add_text(doc, "실제 초도량이 ML 권장량보다 적고 실제 출고율이 100%를 초과했다면, 수요 대비 초도 물량 부족 가능성을 우선 검토합니다.")
    add_bullet(doc, "결품 또는 판매 기회 손실이 있었는지 확인합니다.")
    add_bullet(doc, "ML 권장량을 적용했을 때의 기대 출고율과 정상 발주량 범주를 비교합니다.")
    add_bullet(doc, "다음 유사 상품에서는 초도량을 늘리거나 핵심 센터 배분을 강화하는 근거로 사용합니다.")

    # Checklist and troubleshooting
    add_heading(doc, "5. 발주 확정 전 체크리스트", 1)
    checklist = [
        "선택한 날짜가 실제 발주 마감 업무일과 일치한다.",
        "예약주문 4일치가 모두 존재하고 추세가 비정상적이지 않다.",
        "권장 발주량의 단위가 박스인지 EA인지 구분했다.",
        "입력 박스 수 아래의 EA 환산값을 확인했다.",
        "예상 출고율의 색상과 판정 문구를 확인했다.",
        "권장량에서 벗어난 경우 현업 근거가 있다.",
        "센터별 분배량 합과 전체 입력량이 일치한다.",
        "특정 센터에 과도한 물량이 집중되지 않았다.",
        "센터 가중치를 수정했다면 변경 이유와 결과를 확인했다.",
        "최종 검토 후 발주 확정 상태가 ‘확정 완료’로 변경되었다.",
    ]
    for item in checklist:
        add_bullet(doc, f"□ {item}")

    add_heading(doc, "자주 발생하는 상황과 해결 방법", 1)
    add_feature_table(doc, [
        ("금일 탭에 상품이 없음", "선택 날짜에 마감되는 상품이 없거나 예약주문 4일 데이터가 불완전", "날짜를 조회 가능 구간 내 다른 날짜로 변경"),
        ("수량을 지우면 다시 채워짐", "이전 초기화 로직의 문제", "현재 버전에서는 빈칸 유지가 가능하므로 새로고침 후 재입력"),
        ("과거 행 상세가 열리지 않음", "실제 초도 발주량이 0", "행 아래 표시된 데이터 사유 확인"),
        ("센터 소화율이 0.0%", "해당 센터에서 7일 순재고 감소가 없거나 재고 흐름이 0으로 집계", "마우스 오버로 7일 소화 물량 확인 후 원천 재고 검토"),
        ("센터별 합과 과거 발주량이 다름", "데이터 갱신 또는 구버전 캐시 가능성", "강력 새로고침 후 최신 데이터 배치 확인"),
        ("입력 후 센터 수량이 바뀜", "전체 입력 박스가 점포 비중×가중치 기준으로 재배분", "센터별 수량과 총합을 다시 검토"),
        ("예상 출고율이 실제 감각과 다름", "현재 지표가 ML 권장량 대비 입력량 비율이기 때문", "사후 실제 출고율과 혼동하지 말고 보조 지표로 사용"),
    ])

    add_heading(doc, "브라우저 사용 팁", 2)
    add_bullet(doc, "상품 행은 한 번 클릭하면 상세가 열리고, 같은 행을 다시 클릭하면 닫힙니다.")
    add_bullet(doc, "센터 성과 행에 마우스를 올리면 7일 소화 물량 툴팁이 즉시 표시됩니다.")
    add_bullet(doc, "데이터가 갱신되었는데 숫자가 그대로라면 브라우저 강력 새로고침을 수행합니다.")
    add_bullet(doc, "공유 URL을 사용할 때는 Vercel의 Deployment Protection 설정에 따라 로그인 또는 권한 요청이 필요할 수 있습니다.")

    # Data and limitations
    add_heading(doc, "6. 데이터 기준과 해석상의 한계", 1)
    add_heading(doc, "주요 데이터 출처", 2)
    add_feature_table(doc, [
        ("예약주문·센터 발주 원천", "상품, 출시일, 예약주문, 센터 코드, 초기 발주 관련 정보", "금일 대상 선정 및 예약주문 시각화"),
        ("ML 예측 데이터", "ML 권장 초도량 및 관련 예측 값", "권장 발주량과 과거 ML 비교"),
        ("A4_CENTER_STK.csv", "일자·센터·상품별 장부 마감 재고", "실제 초도 발주량, 7일 출고량, 센터 성과 계산"),
        ("상품 분류 정보", "대분류·중분류·소분류", "과거 상품 필터와 유사 상품 조회"),
        ("센터 점포·가중치 정보", "센터별 점포 비중과 가중치", "센터별 권장 분배량 산정"),
    ])

    add_heading(doc, "현재 계산 정의", 2)
    add_bullet(doc, "실제 초도 발주량은 ‘출시일까지 센터 재고의 양(+)의 증가분 합’으로 정의합니다.")
    add_bullet(doc, "과거 발주량은 센터별 실제 초도 분배량의 합과 동일합니다.")
    add_bullet(doc, "7일 출고율은 ‘출시 후 7일 순출고량 ÷ 실제 초도 발주량 × 100’입니다.")
    add_bullet(doc, "센터별 성과는 각 센터의 초도 분배량과 7일 소화량을 기준으로 판정합니다.")
    add_bullet(doc, "ML 권장량이 없는 상품은 ML 비교가 제한되거나 비교 불가로 표시될 수 있습니다.")

    add_heading(doc, "반드시 알아둘 한계", 2)
    add_callout(
        doc,
        "1. 금일 예상 출고율은 현재 통제 지표",
        "현재 구현은 입력 발주량과 ML 권장량의 상대 비율을 사용합니다. 실제 판매·출고를 독립적으로 예측하는 완성형 모델 지표가 아니므로, "
        "향후 predictedOutflow7d 등 별도 예측 출고량을 안정적으로 확보하면 계산 로직을 고도화할 수 있습니다.",
        fill=RED_BG,
        title_color=RED,
    )
    add_callout(
        doc,
        "2. 재고 감소는 출고를 근사",
        "센터 장부 재고의 감소를 출고량으로 해석합니다. 추가 입고, 재고 조정, 반품, 폐기 등 재고 변동 요인이 있으면 실제 물류 출고와 차이가 생길 수 있습니다.",
        fill=AMBER_BG,
        title_color=AMBER,
    )
    add_callout(
        doc,
        "3. 최종 의사결정은 현업 맥락과 결합",
        "대시보드에는 프로모션 강도, 진열 계획, 경쟁 상품, 날씨, 공급 제한 등 모든 요인이 포함되지 않을 수 있습니다. "
        "경고 색상만으로 수량을 자동 결정하지 말고 현업 정보를 함께 검토합니다.",
        fill=BLUE_BG,
        title_color=BLUE,
    )

    # Quick reference
    add_heading(doc, "7. 한 페이지 빠른 참고표", 1)
    add_heading(doc, "금일 작업", 2)
    for text in (
        "날짜 선택 → 대상 상품 확인",
        "검토 필요 상품 우선 확인",
        "상품 행 클릭 → 예약주문·유사상품·센터 분배 확인",
        "발주 수량은 박스로 입력 → EA 환산 확인",
        "센터 가중치 또는 센터별 수량 조정",
        "전체 수량·센터 합계·판정 확인 → 발주 확정",
    ):
        add_number(doc, text)

    add_heading(doc, "과거 분석", 2)
    for text in (
        "날짜와 대·중·소분류 설정",
        "출고율 또는 발주량 기준 정렬",
        "과거 발주량 색상으로 위험 상품 선별",
        "행 클릭 → 예약주문·7일 출고·센터 성과 확인",
        "ML 권장과 실제 초도 편차 해석",
        "향후 유사 상품 발주 근거로 기록 또는 엑셀 다운로드",
    ):
        add_number(doc, text)

    add_heading(doc, "판정 기준", 2)
    add_status_table(doc)

    add_heading(doc, "마지막으로 기억할 세 가지", 2)
    add_callout(doc, "첫째", "입력 단위는 박스, 비교와 계산은 EA 환산값까지 함께 확인합니다.")
    add_callout(doc, "둘째", "금일 예상 출고율과 과거 실제 출고율은 목적과 계산 방식이 다릅니다.")
    add_callout(doc, "셋째", "권장값은 의사결정의 기준점이며, 예약주문·센터 성과·현업 정보를 함께 보고 확정합니다.")

    # Metadata
    props = doc.core_properties
    props.title = "MD 신상품 초도발주 대시보드 사용 매뉴얼"
    props.subject = "금일 신상품 작업 및 과거 신상품 조회 탭 사용 안내"
    props.author = "Seven-Eleven MD Dashboard Project"
    props.keywords = "MD, 신상품, 초도발주, 대시보드, 매뉴얼"
    props.comments = "현재 React 대시보드 구현 및 데이터 산출 로직 기준"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
